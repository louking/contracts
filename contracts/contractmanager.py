###########################################################################################
# contracts - contracts views
#
#       Date            Author          Reason
#       ----            ------          ------
#       10/13/18        Lou King        Create
#
#   Copyright 2018 Lou King
#
###########################################################################################
'''
contract - class and helpers to manage contract
==================================================
'''

# standard
from tempfile import mkdtemp
from datetime import date
from os.path import join as pathjoin
from copy import deepcopy
from csv import reader
from shutil import rmtree
from json import loads
from os.path import join as pathjoin

# pypy
from docx import Document
from flask import current_app, redirect
from jinja2 import Environment
from slugify import slugify
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload

# homegrown
from contracts.dbmodel import db, Contract, ContractType, ContractBlockType, TemplateType
from loutilities import timeu
from loutilities.googleauth import GoogleAuthService
from .html2docx import convert

class parameterError(Exception): pass

debug = True
debug2 = False

# templating environment. strip white space. see http://jinja.pocoo.org/docs/2.10/templates/#whitespace-control
template_env = Environment(trim_blocks=True, lstrip_blocks=True)

# needful constants
DRIVE_SERVICE = 'drive'
DRIVE_VERSION = 'v3'

# exceptions
class PermissionError(Exception): pass

#----------------------------------------------------------------------
def _evaluate(tree, subtree):
#----------------------------------------------------------------------
    '''
    detect callable attributes of the subtree object, and evaluate them
    by replacing with result of subtree.attr(tree)
    '''
    # adapted from https://stackoverflow.com/questions/7963762/what-is-the-most-economical-way-to-convert-nested-python-objects-to-dictionaries
    # if at a leaf, either return the leaf of leaf(tree)
    if not hasattr(subtree, '__dict__'):
        if not callable(subtree):
            return subtree
        else:
            return subtree(tree)

    # flow here if we have an object
    for key, val in list(subtree.__dict__.items()):
        if key.startswith('_'): continue

        element = []
        if isinstance(val, list):
            for item in val:
                if debug2: current_app.logger.debug('_evaluate(): processing list item key={} item={}'.format(key, item))
                element.append( _evaluate( tree, item ) )
        else:
            element = _evaluate( tree, getattr(subtree, key) )

        if debug2: current_app.logger.debug('_evaluate(): key={} element={}'.format(key, element))
        setattr( subtree, key, element )

    # bubble up the results
    return subtree

#----------------------------------------------------------------------
def recursive_render(tpl, values):
#----------------------------------------------------------------------
    '''
    recursively render jinja2 template 
    from https://stackoverflow.com/questions/8862731/jinja-nested-rendering-on-variable-content

    :param tpl: template text
    :param values: values to be rendered
    '''
    # create environment
    env = Environment(trim_blocks=True, lstrip_blocks=True, loader=current_app.jinja_env.loader)

### TODO: note can only be used by html, for docx need to consider generator aspect within ContractManagerTemplate
    prev = tpl
    while True:
        curr = env.from_string(prev).render(**values)
        if curr != prev:
            prev = curr
        else:
            return curr

#####################################################
class ContractManagerTemplate():
#####################################################
    '''
    parameters:

    * template - jinja2 template of text with replacement fields surrounded by curly braces with fields 
      like {{ a }, {{ b.c }} and control like {% for xxx %} {% endfor %}, {% if xxx %} {% endif %} 
      see http://jinja.pocoo.org/docs/2.10/templates/
    '''

    #----------------------------------------------------------------------
    def __init__(self, template):
    #----------------------------------------------------------------------
        if debug: current_app.logger.debug('ContractManager.__init__(): template={}'.format(template))
        self.template = template_env.from_string(template)

    #----------------------------------------------------------------------
    def render(self, mergefields):
    #----------------------------------------------------------------------
        '''
        merge a template of document output based on mergefields

        parameters:

        * merge - object containing items to be substituted into format strings, may have callable items  
          top level is turned into dict and subsequent levels are treated as attributes
        '''
        # preprocess any callable leaves
        merge = _evaluate(mergefields, mergefields)

        # render the template
        if debug: current_app.logger.debug('ContractManager.render(): merge={}'.format(merge.__dict__))
        evaluated = self.template.render(**merge.__dict__)

        return evaluated
    
    #----------------------------------------------------------------------
    def generate(self, mergefields):
    #----------------------------------------------------------------------
        '''
        create jinja2 generator to render text incrementally.

        parameters:

        * takes same arguments as render

        returns: generator object, generator.next() yields next section or raises StopIteration
        '''
        # preprocess any callable leaves
        merge = _evaluate(mergefields, mergefields)

        # create the template generator
        if debug: current_app.logger.debug('ContractManager.generate(): merge={}'.format(merge.__dict__))
        generator = self.template.generate(**merge.__dict__)

        return generator

#####################################################
class ContractManager():
#####################################################
    '''
    manage contract creation, storage and retrieval

    parameters:

    * contracttype - type of contract
    * driveFolder - folder id on google drive to store contract document
    '''

    #----------------------------------------------------------------------
    def __init__(self, **kwargs):
    #----------------------------------------------------------------------
        # the args dict has default values for arguments added by this class
        # caller supplied keyword args are used to update these
        # all arguments are made into attributes for self by the inherited class
        args = dict(contractType=None,
                    templateType=None,
                    driveFolderId=None,
                    doctype='docx', # can be docx or html
                    )
        args.update(kwargs)
        for key in args:
            setattr(self, key, args[key])

        if self.doctype not in ['docx', 'html']:
            raise parameterError('ContractManager(): doctype must be "docx" or "html", found {}'.format(self.doctype))

    #----------------------------------------------------------------------
    def create(self, filename, mergefields, addlfields={}, is_quote=True):
    #----------------------------------------------------------------------
        '''
        create the document

        parameters:

        * filename - name of file to be created
          * note different processing for .docx vs. .html
        * mergefields - flat dict with keys to be used as merge fields. If field contains 
          a function, the function is called with a single argument: the original mergefields itself
        * addlfields - any fields to be added to mergefields
        * is_quote - True if quote being created, False if invoice being created

        returns: G Suite document id
        '''

        # create document based on doctype
        if self.doctype == 'docx':
            # docx handle
            docx = Document()
        elif self.doctype == 'html':
            html = []
        else:
            raise parameterError('create(): bad doctype {}'.format(self.doctype))


        # retrieve contract template
        templates = (db.session.query(Contract)
                     .filter(Contract.contractTypeId==ContractType.id)
                     .filter(ContractType.contractType==self.contractType)
                     .filter(Contract.templateTypeId==TemplateType.id)
                     .filter(TemplateType.templateType==self.templateType)
                     .order_by(Contract.blockPriority)
                     .all()
                    )

        # prepare built in fields
        dt = timeu.asctime('%B %d, %Y')
        _date_ = dt.dt2asc( date.today() )
        
        # copy caller's mergefields and add built-in fields
        course =  getattr(mergefields, 'course', None)  # otherwise lazy load doesn't work
        merge = deepcopy(mergefields)
        merge._date_ = _date_

        # maybe there are some additional fields
        for key in addlfields:
            setattr(merge, key, addlfields[key])

        # fill contents for html files
        if self.doctype == 'html':
            for blockd in templates:
                # retrieve block type text
                blockType = blockd.contractBlockType.blockType

                # only possibility is html
                if blockType == 'html':
                    html.append( recursive_render( blockd.block, merge.__dict__ ) )
                    # template = ContractManagerTemplate( blockd.block )
                    # html.append( template.render( merge ) )
                
                # bad configuration, or bad code
                else:
                    raise parameterError('unknown block type for {}: {}'.format(self.doctype, blockType))

        # fill contents for docx files
        elif self.doctype == 'docx':
            # invoicestart and invoiceend are special ContractBlockTypes for docx doctype
            invoicestart = db.session.query(ContractBlockType).filter_by(blockType='invoicestart').one_or_none()
            invoiceend = db.session.query(ContractBlockType).filter_by(blockType='invoiceend').one_or_none()

            # for contract quote or sponsor agreement, remove 'invoicestart' and 'invoiceend'
            if is_quote:
                templates = [t for t in templates if t.contractBlockType != invoicestart and t.contractBlockType != invoiceend]
            
            # for invoice, remove all templates outside of 'invoicestart' and 'invoiceend'
            else:
                if invoicestart:
                    while len(templates) > 0 and templates[0].contractBlockType != invoicestart:
                        templates.pop(0)
                    if len(templates) > 0 and templates[0].contractBlockType == invoicestart:
                        templates.pop(0)
                if invoiceend:
                    while len(templates) > 0 and templates[-1].contractBlockType != invoiceend:
                        templates.pop()
                    if len(templates) > 0 and templates[-1].contractBlockType == invoiceend:
                        templates.pop()
                
            for blockd in templates:
                # retrieve block type text
                blockType = blockd.contractBlockType.blockType

                # para is a single paragraph, based on a single template
                if blockType == 'para':
                    template = ContractManagerTemplate( blockd.block )
                    para = docx.add_paragraph( template.render( merge ) )

                # sectionprops has json object containing section property assignments (see https://python-docx.readthedocs.io/en/latest/api/section.html)
                elif blockType == 'sectionprops':
                    props = loads( blockd.block )
                    for prop in props:
                        # try to convert property value, otherwise leave as string
                        thisprop = props[prop]
                        for ttype in [int, float]:
                            try:
                                thisprop = ttype(thisprop)
                                break
                            except ValueError:
                                pass
                        setattr(docx.sections[0], prop, thisprop)

                # listitem[2] is a list item which may generate multiple lines, based on a single template
                elif blockType in ['listitem', 'listitem2']:
                    template = ContractManagerTemplate( blockd.block )
                    for render in template.generate( merge ):
                        listitem = docx.add_paragraph( render )
                        if blockType == 'listitem':
                            listitem.style = 'List Bullet'
                        elif blockType == 'listitem2':
                            listitem.style = 'List Bullet 2'
                    
                # pagebreak is, well, just a page break
                elif blockType == 'pagebreak':
                    docx.add_page_break()
                    
                # tablehdr causes a table to be created with the number of columns depending on how many elements
                # this is configured as comma separated headings (no template rendering is done)
                elif blockType == 'tablehdr':
                    # use csv reader to parse quoted fields with commas correctly
                    rdr = reader([blockd.block])
                    headings = next(rdr)
                    table = docx.add_table(rows=1, cols=len(headings))
                    hdr_cells = table.rows[0].cells
                    for c in range(len(headings)):
                        run = hdr_cells[c].paragraphs[0].add_run( headings[c] )
                        run.bold = True

                # tablerow and tablerowbold define what some rows of the table look like
                # this is configured as template with comma separated columns, optional for loop
                elif blockType in ['tablerow', 'tablerowbold']:
                    # get templated and create generator
                    coltemplate = ContractManagerTemplate( blockd.block )
                    colg = coltemplate.generate( merge )

                    # collect the generated raw rows (raw means with commas embedded)
                    rawrows = []
                    for rawrow in colg:
                        rawrows.append( rawrow )

                    # then add row data to table using csv reader 
                    # if there are commas in the data csv reader will parse quoted fields with commas correctly
                    rows = reader( rawrows )
                    for row in rows:
                        row_cells = table.add_row().cells
                        for c in range(len(row)):
                            run = row_cells[c].paragraphs[0].add_run( row[c] )
                            if blockType == 'tablerowbold':
                                run.bold = True
                
                else:
                    raise parameterError('unknown block type for {}: {}'.format(self.doctype, blockType))

        # save temporary doc file
        dirpath = mkdtemp(prefix='contracts_')

        # save file, set up for drive upload
        drivename = filename

        # for mimetype see https://developer.mozilla.org/en-US/docs/Web/HTTP/Basics_of_HTTP/MIME_types/Complete_list_of_MIME_types
        if self.doctype == 'html':
            if drivename[-5:] == '.html': drivename = drivename[:-5]
            htmlpath = pathjoin(dirpath, slugify(drivename)) + '.html'
            with open(htmlpath,'w') as htmlfile:
                htmlfile.write('\n'.join(html))
            # converted docx file gets uploaded to drive
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            # mimetype='text/html'
            path = pathjoin(dirpath, slugify(drivename)) + '.docx'
            convert(htmlpath, path, title=drivename)
        elif self.doctype == 'docx':
            if drivename[-5:] == '.docx': drivename = drivename[:-5]
            # slugify to avoid funky characters in race name
            path = pathjoin(dirpath, slugify(drivename)) + '.docx'
            docx.save(path)
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        if debug: current_app.logger.debug('ContractManager.create(): created temporary {}'.format(path))

        # upload to google drive
        gs = GoogleAuthService(current_app.config['GSUITE_SERVICE_KEY_FILE'], current_app.config['GSUITE_SCOPES'])
        fid = gs.create_file(current_app.config['CONTRACTS_DB_FOLDER'], drivename, path, doctype='docx')
        if debug: current_app.logger.debug('uploaded fid={}'.format(fid))

        ## set file to be publicly readable
        public_permission = {
            'type': 'anyone',
            'role': 'reader',
        }
        gs.set_permission(fid, public_permission)   

        # remove temporary folder
        # NOTE: in windows at least, this gets an error because file is still in use
        try:
            rmtree(dirpath, ignore_errors=True)
        except:
            pass

        # send email

        return fid


