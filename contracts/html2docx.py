# adapted from https://github.com/erezlife/html2docx

import io
import re

# original software built for python 3.6+
import sys
python2 = sys.version_info[0] == 2
if python2:
    from HTMLParser import HTMLParser
else:
    from html.parser import HTMLParser

import docx

from docx.shared import Inches, Cm, Mm, Pt, Twips, Emu
csslen2docx = {
            'px': lambda val: Pt(96.0*val/72.0), 
            'cm': lambda val: Cm(val), 
            'mm': lambda val: Mm(val), 
            'in': lambda val: Inches(val), 
            'pc': lambda val: Inches(val/6.0),
            'pt': lambda val: Pt(val),
           }
LEN_RE = re.compile(r"([0-9]*\.?[0-9]+)([a-z]+)")

WHITESPACE_RE = re.compile(r"\s+")

# support python 2.7
import fmt as f

class HTML2Docx(HTMLParser):
    header_re = re.compile(r"^h[1-6]$")

    def __init__(self, title):
        HTMLParser.__init__(self)
        self.doc = docx.Document()
        self.doc.core_properties.title = title

        self.p = None
        self.header = None
        self.footer = None
        self.runs = []
        self.list_style = []
        self.bold = 0
        self.italic = 0
        self.underline = 0

    def feed(self, data):
        HTMLParser.feed(self,data)
        self.finish()

    def handle_starttag(self, tag, attrs):
        if tag in ["p", "div"]:
            self.finish()
            self.p = self.doc.add_paragraph()
        elif tag == "br":
            self.out("\n", strip_whitespace=False)
        elif self.header_re.match(tag):
            self.finish()
            level = int(tag[-1])
            self.p = self.doc.add_heading(level=level)
        elif tag == "ul":
            self.add_list_style("List Bullet")
        elif tag == "ol":
            self.add_list_style("List Number")
        elif tag == "li":
            self.p = self.doc.add_paragraph(style=self.list_style[-1])
        elif tag in ["b", "strong"]:
            self.bold += 1
        elif tag in ["i", "em"]:
            self.italic += 1
        elif tag == "u":
            self.underline += 1
        elif tag == "img":
            dattrs = dict(attrs)
            width = dattrs.get('width')
            height = dattrs.get('height')
            if width:
                m = LEN_RE.match(dattrs['width'])
                width = csslen2docx[m.group(2)](float(m.group(1)))
            if height:
                m = LEN_RE.match(dattrs['height'])
                width = csslen2docx[m.group(2)](float(m.group(1)))
            if self.header:
                run = self.p.add_run()
                run.add_picture(dattrs['src'], width=width, height=height)
            elif self.footer:
                run = self.p.add_run()
                run.add_picture(dattrs['src'], width=width, height=height)
            else:
                self.doc.add_picture(dattrs['src'], width=width, height=height)
        
        # non-standard tag for header, footer - these act like p tag
        elif tag == "header":
            self.finish()
            section = self.doc.sections[0]
            self.header = section.header
            self.p = self.header.paragraphs[0]
        elif tag == "footer":
            self.finish()
            section = self.doc.sections[0]
            self.footer = section.footer
            self.p = self.footer.paragraphs[0]

    def handle_endtag(self, tag):
        if tag in ["p", "div", "li"] or self.header_re.match(tag):
            self.finish()
            self.p = None
        elif tag in ["ul", "ol"]:
            self.list_style.pop()
        elif tag in ["b", "strong"]:
            self.bold -= 1
        elif tag in ["i", "em"]:
            self.italic -= 1
        elif tag == "u":
            self.underline -= 1

        # non-standard tag for header, footer - these act like p tag
        elif tag == "header":
            self.finish()
            self.p = None
            self.header = None
        elif tag == "footer":
            self.finish()
            self.p = None
            self.footer = None

    def handle_data(self, data):
        if self.p or data.strip():
            self.out(data)

    def out(self, data, strip_whitespace=True):
        if not self.p:
            self.p = self.doc.add_paragraph()

        run = Run(
            data,
            bold=bool(self.bold),
            italic=bool(self.italic),
            underline=bool(self.underline),
            strip_whitespace=strip_whitespace,
        )
        self.runs.append(run)

    def finish(self):
        """Collapse white space across runs."""
        prev_run = None
        for run in self.runs:
            if run.strip_whitespace:
                run.text = WHITESPACE_RE.sub(" ", run.text.strip())
                if prev_run and prev_run.strip_whitespace:
                    if prev_run.needs_space_suffix:
                        prev_run.text += " "
                    elif run.text and run.needs_space_prefix:
                        run.text = " " + run.text
            prev_run = run

        for run in self.runs:
            if run.text:
                doc_run = self.p.add_run(run.text)
                doc_run.bold = run.bold
                doc_run.italic = run.italic
                doc_run.underline = run.underline

        self.p = None
        self.runs = []

    def add_list_style(self, style):
        self.finish()
        if self.list_style:
            # The template included by python-docx only has 3 list styles.
            level = min(len(self.list_style) + 1, 3)
            style = f("{style} {level}")
            # requires python 3.6+
            # style = f"{style} {level}"
        self.list_style.append(style)


class Run:
    def __init__(
        self, text, bold=False, italic=False, underline=False, strip_whitespace=True
    ):
        self.text = text
        self.bold = bold
        self.italic = italic
        self.underline = underline
        self.strip_whitespace = strip_whitespace

        self.needs_space_prefix = bool(WHITESPACE_RE.match(text[0]))
        self.needs_space_suffix = bool(WHITESPACE_RE.match(text[-1]))


def html2docx(content, title):
    """Convert valid HTML content to a docx document and return it as a
    io.BytesIO() object.

    """
    parser = HTML2Docx(title)
    parser.feed(content)

    buf = io.BytesIO()
    parser.doc.save(buf)
    return buf

def convert(htmlname, docxname, title=None):
    with open(htmlname) as fp:
        html = fp.read()
    
    buf = html2docx(html, title=title)
    with open(docxname, 'wb') as fp:
        fp.write(buf.getvalue())