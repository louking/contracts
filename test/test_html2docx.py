'''
test_html2docx - test contracts.html2docx
=========================================================
'''

# standard
import os

# pypi
from docx import Document

# homegrown
from contracts.html2docx import html2docx


def _docx(html):
    buf = html2docx(html, title='test doc')
    return Document(buf)


def test_paragraph_text():
    doc = _docx('<p>Hello world</p>')
    assert [p.text for p in doc.paragraphs] == ['Hello world']


def test_title_set_on_core_properties():
    doc = _docx('<p>Hello world</p>')
    assert doc.core_properties.title == 'test doc'


def test_bold_and_italic_runs():
    doc = _docx('<p>Hello <b>bold</b> and <i>italic</i></p>')
    para = doc.paragraphs[0]
    runs = [(r.text, r.bold, r.italic) for r in para.runs]
    assert ('Hello ', False, False) in runs
    assert ('bold', True, False) in runs
    assert (' and ', False, False) in runs
    assert ('italic', False, True) in runs


def test_underline_run():
    doc = _docx('<p><u>underlined</u></p>')
    para = doc.paragraphs[0]
    assert para.runs[0].underline is True


def test_heading_levels():
    doc = _docx('<h1>Big</h1><h2>Smaller</h2>')
    assert doc.paragraphs[0].style.name == 'Heading 1'
    assert doc.paragraphs[0].text == 'Big'
    assert doc.paragraphs[1].style.name == 'Heading 2'
    assert doc.paragraphs[1].text == 'Smaller'


def test_unordered_list_items_use_bullet_style():
    doc = _docx('<ul><li>one</li><li>two</li></ul>')
    assert [p.text for p in doc.paragraphs] == ['one', 'two']
    assert all(p.style.name == 'List Bullet' for p in doc.paragraphs)


def test_ordered_list_items_use_number_style():
    doc = _docx('<ol><li>first</li><li>second</li></ol>')
    assert [p.text for p in doc.paragraphs] == ['first', 'second']
    assert all(p.style.name == 'List Number' for p in doc.paragraphs)


def test_nested_list_uses_level_2_style():
    doc = _docx('<ul><li>top<ul><li>nested</li></ul></li></ul>')
    styles = {p.text: p.style.name for p in doc.paragraphs}
    assert styles['top'] == 'List Bullet'
    assert styles['nested'] == 'List Bullet 2'


def test_multiple_paragraphs():
    doc = _docx('<p>First</p><p>Second</p>')
    assert [p.text for p in doc.paragraphs] == ['First', 'Second']


def test_br_inserts_newline_within_paragraph():
    doc = _docx('<p>line one<br>line two</p>')
    assert doc.paragraphs[0].text == 'line one\nline two'


def test_whitespace_is_collapsed():
    doc = _docx('<p>  lots   of    space  </p>')
    assert doc.paragraphs[0].text == 'lots of space'


def test_div_treated_like_paragraph():
    doc = _docx('<div>divtext</div>')
    assert [p.text for p in doc.paragraphs] == ['divtext']
